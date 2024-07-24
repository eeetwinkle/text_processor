import sys
import re
from PyQt5 import QtWidgets, QtGui, QtCore
from QtMainWindow import Ui_color
from QtSearchWindow import Ui_QtSearchWindow
from QtReplaceWindow import Ui_QtReplaceWindow
from QtStyles import Ui_Form
from QtNewStyle import Ui_QtNewStyleWindow
from PyQt5.QtGui import QTextCursor, QTextBlockFormat, QTextImageFormat, QPixmap
from PyQt5.QtWidgets import QColorDialog, QFileDialog, QMessageBox, QInputDialog
from docx import Document
from docx.shared import Pt, RGBColor


class MainWindow(QtWidgets.QMainWindow, Ui_color):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.search_window = SearchWindow(self.text_edit)
        self.replace_window = ReplaceWindow(self.text_edit)
        self.style_window = StyleWindow()
        self.search.clicked.connect(self.open_search_window)
        self.replace.clicked.connect(self.open_replace_window)
        self.style.clicked.connect(self.open_style_window)
        self.text_color.clicked.connect(self.change_text_color)
        self.save.clicked.connect(self.save_as_docx)

        self.bold_active = False
        self.italic_active = False
        self.underlined_active = False

        self.font.currentFontChanged.connect(self.update_font)
        self.font_size.currentIndexChanged.connect(self.update_font_size)

        self.bold.clicked.connect(self.toggle_bold)
        self.italic.clicked.connect(self.toggle_italic)
        self.underlined.clicked.connect(self.toggle_underlined)

        self.update_font()
        self.update_font_size()

        self.indent_value = 0
        self.reduce_indentation.clicked.connect(self.update_indent)
        self.increase_indentation.clicked.connect(self.update_indent)

        self.paste.clicked.connect(self.insert_image)

    def update_indent(self):
        sender = self.sender()
        if sender == self.increase_indentation:
            self.indent_value += 1
        elif self.indent_value > 0:
            self.indent_value -= 1
        block_format = QTextBlockFormat()
        block_format.setIndent(self.indent_value)
        cursor = self.text_edit.textCursor()
        cursor.select(QTextCursor.Document)
        cursor.mergeBlockFormat(block_format)
        self.text_edit.setTextCursor(cursor)

    def update_font(self, font=None):
        if font is None:
            font = self.font.currentFont()
        self.apply_font(font)

    def update_font_size(self):
        size = int(self.font_size.currentText())
        self.apply_font_size(size)

    def apply_font(self, font):
        format = QtGui.QTextCharFormat()
        format.setFontFamily(font.family())
        self.merge_format_on_word_or_selection(format)

    def apply_font_size(self, size):
        format = QtGui.QTextCharFormat()
        format.setFontPointSize(size)
        self.merge_format_on_word_or_selection(format)

    def toggle_bold(self):
        self.bold_active = not self.bold_active
        self.bold.setStyleSheet('background-color: lightblue' if self.bold_active else 'background-color: none')
        format = QtGui.QTextCharFormat()
        format.setFontWeight(QtGui.QFont.Bold if self.bold_active else QtGui.QFont.Normal)
        self.merge_format_on_word_or_selection(format)

    def toggle_italic(self):
        self.italic_active = not self.italic_active
        self.italic.setStyleSheet('background-color: lightblue' if self.italic_active else 'background-color: none')
        format = QtGui.QTextCharFormat()
        format.setFontItalic(self.italic_active)
        self.merge_format_on_word_or_selection(format)

    def toggle_underlined(self):
        self.underlined_active = not self.underlined_active
        self.underlined.setStyleSheet(
            'background-color: lightblue' if self.underlined_active else 'background-color: none')
        format = QtGui.QTextCharFormat()
        format.setFontUnderline(self.underlined_active)
        self.merge_format_on_word_or_selection(format)
    def change_text_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            format = QtGui.QTextCharFormat()
            format.setForeground(color)
            self.merge_format_on_word_or_selection(format)
            self.text_color.setStyleSheet(
                f'background-color: {color.name()}' if color != QtGui.QColor('black') else 'background-color: none')

    def merge_format_on_word_or_selection(self, format):
        cursor = self.text_edit.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        cursor.mergeCharFormat(format)
        self.text_edit.mergeCurrentCharFormat(format)

    def update_current_format(self):
        format = self.text_edit.currentCharFormat()
        format.setFontWeight(QtGui.QFont.Bold if self.bold_active else QtGui.QFont.Normal)
        format.setFontItalic(self.italic_active)
        format.setFontUnderline(self.underlined_active)
        self.text_edit.setCurrentCharFormat(format)

    def save_as_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Word Documents (*.docx)")
        if file_path:
            doc = Document()
            cursor = self.text_edit.textCursor()
            cursor.select(QTextCursor.Document)
            text = cursor.selection().toPlainText()

            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(para)
                    format = cursor.charFormat()

                    font = run.font
                    font.name = format.fontFamily()
                    font.size = Pt(format.fontPointSize())
                    font.bold = format.fontWeight() == QtGui.QFont.Bold
                    font.italic = format.fontItalic()
                    font.underline = format.fontUnderline()

                    color = format.foreground().color()
                    font.color.rgb = RGBColor(color.red(), color.green(), color.blue())

            doc.save(file_path)

    def insert_image(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Выберите изображение", "",
                                                   "Images (*.png *.xpm *.jpg *.jpeg);;All Files (*)", options=options)
        if file_name:
            try:
                cursor = self.text_edit.textCursor()
                image_format = QTextImageFormat()
                image_format.setName(file_name)

                # Получаем исходные размеры изображения
                pixmap = QPixmap(file_name)
                if pixmap.isNull():
                    raise ValueError("Изображение не может быть загружено.")

                # Запрашиваем новые размеры у пользователя с предложением использовать исходные
                default_width = pixmap.width()
                default_height = pixmap.height()
                width, ok = QInputDialog.getInt(self, "Ширина изображения", "Введите ширину:", default_width, 1, 3000)
                if ok:
                    height, ok = QInputDialog.getInt(self, "Высота изображения", "Введите высоту:", default_height, 1,
                                                     3000)
                    if ok:
                        image_format.setWidth(width)
                        image_format.setHeight(height)
                        cursor.insertImage(image_format)
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось вставить изображение: {str(e)}")

    def open_search_window(self):
        self.search_window.show()

    def open_replace_window(self):
        self.replace_window.show()

    def open_style_window(self):
        self.style_window.show()


class SearchWindow(QtWidgets.QWidget, Ui_QtSearchWindow):
    def __init__(self, text_edit):
        super().__init__()
        self.text_edit = text_edit
        self.setupUi(self)
        self.pushButton_search.clicked.connect(self.perform_search)
        self.pushButton_search_2.clicked.connect(self.navigate_to_next)
        self.pushButton_search_3.clicked.connect(self.navigate_to_previous)

    # Метод, выполняющий поиск по тексту
    def perform_search(self):
        search_text = self.lineEdit_search.text()
        if self.checkBox_register.isChecked():
            flags = 0
        else:
            flags = re.IGNORECASE  # Учитываем регистр иначе

        whole_word = self.checkBox_entirely.isChecked()
        pattern = search_text
        if whole_word:
            pattern = r'\b' + re.escape(search_text) + r'\b'

        text = self.text_edit.toPlainText()
        self.found_positions = [m.start() for m in re.finditer(pattern, text, flags)]
        count = len(self.found_positions)

        if count == 0:
            self.show_message("Не найдено")
            self.current_index = -1  # Сбрасываем индекс
            self.clear_highlight()
        else:
            self.show_message(f"Найдено: {count} вхождений")
            self.current_index = 0  # Сбрасываем индекс для первого вхождения
            self.highlight_current_word()

    # Метод для выделения текущего найденного слова
    def highlight_current_word(self):
        self.clear_highlight()  # Сначала снимаем выделение

        if self.current_index >= 0 and self.current_index < len(self.found_positions):
            cursor = self.text_edit.textCursor()
            pos = self.found_positions[self.current_index]
            cursor.setPosition(pos)
            cursor.movePosition(cursor.Right, cursor.KeepAnchor, len(self.lineEdit_search.text()))
            self.text_edit.setTextCursor(cursor)
            self.text_edit.setFocus()

    # Метод для снятия выделения
    def clear_highlight(self):
        cursor = self.text_edit.textCursor()
        cursor.clearSelection()
        self.text_edit.setTextCursor(cursor)

    # Метод для отображения сообщения в диалоговом окне
    def show_message(self, message):
        msg_box = QMessageBox()
        msg_box.setText(message)
        msg_box.exec_()

    # Навигация к следующему найденному слову
    def navigate_to_next(self):
        if not self.found_positions:
            return
        self.current_index = (self.current_index + 1) % len(self.found_positions)
        self.highlight_current_word()

    # Навигация к предыдущему найденному слову
    def navigate_to_previous(self):
        if not self.found_positions:
            return
        self.current_index = (self.current_index - 1) % len(self.found_positions)
        self.highlight_current_word()



class ReplaceWindow(QtWidgets.QWidget, Ui_QtReplaceWindow):
    def __init__(self, text_edit):
        super().__init__()
        self.text_edit = text_edit
        self.setupUi(self)
        self.pushButton_replace.clicked.connect(self.replace_all)

    # Метод, выполняющий замену всех найденных слов
    def replace_all(self):
        word_to_replace = self.lineEdit_search2.text()
        replacement_word = self.lineEdit_replace.text()

        if not word_to_replace:
            self.show_message("Пожалуйста, введите слово для замены.")
            return

        if self.checkBox_register.isChecked():
            flags = 0
        else:
            flags = re.IGNORECASE  # Учитываем регистр иначе

        whole_word = self.checkBox_entirely.isChecked()
        pattern = word_to_replace
        if whole_word:
            pattern = r'\b' + re.escape(word_to_replace) + r'\b'

        text = self.text_edit.toPlainText()
        new_text = re.sub(pattern, replacement_word, text, flags=flags)
        self.text_edit.setPlainText(new_text)

        self.show_message(f"Все вхождения '{word_to_replace}' заменены на '{replacement_word}'.")

    # Метод для отображения сообщения в диалоговом окне
    def show_message(self, message):
        msg_box = QMessageBox()
        msg_box.setText(message)
        msg_box.exec_()


class StyleWindow(QtWidgets.QWidget, Ui_Form):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.new_style_window = NewStyleWindow()
        self.new_style.clicked.connect(self.open_new_style_window)

    def open_new_style_window(self):
        self.new_style_window.show()


class NewStyleWindow(QtWidgets.QWidget, Ui_QtNewStyleWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    main_window = MainWindow()
    main_window.show()
    sys.exit(app.exec_())